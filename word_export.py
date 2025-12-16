from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def export_program_to_word(program, carrier_data=None):
    doc = Document()

    # Title
    title = doc.add_heading(program["account"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Program Structure", level=1)

    # Sort layers by attachment
    layers_sorted = sorted(program["layers"], key=lambda x: x.get("attachment", 0))

    for idx, layer in enumerate(layers_sorted):
        limit_val = layer.get("limit", 0)
        attach_val = layer.get("attachment", 0)

        if layer.get("is_primary"):
            layer_title = f"Layer {idx+1}: ${limit_val:,.0f} Primary"
        else:
            layer_title = f"Layer {idx+1}: ${limit_val:,.0f} xs ${attach_val:,.0f}"

        doc.add_heading(layer_title, level=2)

        # Improvement #2 - Add table for all layers
        carriers = layer.get("carriers", [])
        if carriers:
            # Create carrier table
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"

            # Header row
            hdr_cells = table.rows[0].cells
            headers = [
                "Carrier",
                "Share %",
                "Premium ($)",
                "Policy #",
                "Total Fees ($)",
            ]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add carrier rows
            for carrier in carriers:
                carrier_name = carrier.get("carrier_name", "Unknown")
                carrier_share = carrier.get("share", 0)
                premium = carrier.get("premium", 0)
                policy_num = carrier.get("policy_number", "")
                carrier_fee = carrier.get("carrier_fee", 0)
                surplus_fee = carrier.get("surplus_fee", 0)
                total_fees = carrier_fee + surplus_fee

                # Handle Multiple RBEs policy number display
                if (
                    carrier.get("has_multiple_rbes", False)
                    and not carrier.get("single_policy_number", False)
                    and carrier.get("rbes", [])
                ):
                    policy_num = "Multiple"

                # Add row to table
                row_cells = table.add_row().cells
                row_cells[0].text = carrier_name
                row_cells[1].text = f"{carrier_share*100:.1f}%"
                row_cells[2].text = f"${premium:,.0f}"
                row_cells[3].text = policy_num
                row_cells[4].text = f"${total_fees:,.2f}"

                # If carrier has multiple RBEs, add RBE breakdown
                if carrier.get("has_multiple_rbes", False):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run(
                        f"RBE breakdown of {carrier_name}'s {carrier_share*100:.1f}% layer participation:"
                    ).italic = True

                    # Create RBE table
                    rbe_table = doc.add_table(rows=1, cols=5)
                    rbe_table.style = "Light Grid Accent 1"

                    # RBE header row
                    rbe_hdr_cells = rbe_table.rows[0].cells
                    rbe_headers = [
                        "RBE",
                        "RBE Share %",
                        "Layer Share %",
                        "Premium ($)",
                        "Policy #",
                    ]
                    for i, header in enumerate(rbe_headers):
                        rbe_hdr_cells[i].text = header
                        for paragraph in rbe_hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # RBE data rows
                    for rbe in carrier.get("rbes", []):
                        rbe_name = rbe.get("rbe", "")
                        rbe_share = rbe.get("share", 0)
                        layer_share = rbe_share * carrier_share
                        rbe_premium = rbe.get("premium", 0)

                        # Use carrier policy number if single policy is enabled
                        rbe_policy = rbe.get("policy_number", "")
                        if carrier.get("single_policy_number", False):
                            rbe_policy = carrier.get("policy_number", "")

                        # Add row to RBE table
                        rbe_row_cells = rbe_table.add_row().cells
                        rbe_row_cells[0].text = rbe_name
                        rbe_row_cells[1].text = f"{rbe_share*100:.1f}%"
                        rbe_row_cells[2].text = f"{layer_share*100:.2f}%"
                        rbe_row_cells[3].text = f"${rbe_premium:,.0f}"
                        rbe_row_cells[4].text = rbe_policy

            doc.add_paragraph()  # Space after carrier table
        else:
            doc.add_paragraph("No carriers in this layer")

        doc.add_paragraph()  # Space between layers

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
